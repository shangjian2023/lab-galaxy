#!/usr/bin/env python3
"""生成创新实验知识图谱平台实验报告 (Word格式)."""

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = 'KaiTi'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_font(run, name='KaiTi', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_custom(text, level=1):
    sizes = {1: 22, 2: 16, 3: 14}
    s = sizes.get(level, 12)
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=s, bold=True)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, indent=True, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74 * 2)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table_with_data(headers, data, col_widths=None):
    """添加一个带表头的表格。"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, size=10, bold=True)
    # 数据行
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=10)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('创新实验知识图谱平台')
set_font(run, size=28, bold=True, color=(0, 0, 0))
title_p.paragraph_format.space_after = Pt(12)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('实 验 报 告')
set_font(run, size=22, bold=True, color=(0, 0, 0))
subtitle_p.paragraph_format.space_after = Pt(40)

info_lines = [
    ('课程名称：', '计算机综合实验'),
    ('项目名称：', '创新实验知识图谱平台（LabGalaxy）'),
    ('版本号：', 'v0.1.0'),
    ('提交日期：', datetime.date.today().strftime('%Y年%m月%d日')),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label)
    set_font(r1, size=14, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=14)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 摘要 & 关键词
# ═══════════════════════════════════════════════
add_heading_custom('摘  要', level=1)

add_body(
    '本项目设计并实现了一套基于AI驱动的创新实验知识图谱平台（LabGalaxy），旨在解决传统实验教学中知识碎片化、隐性知识难以显性化、跨实验关联发现困难等核心痛点。平台采用前后端分离的微服务架构，整合了大语言模型（LLM）实体关系抽取、Neo4j图数据库存储、FAISS向量索引检索、PostgreSQL关系型数据管理等核心技术，构建了一个涵盖文档上传与解析、知识自动萃取、图谱可视化探索、自然语言问答、团队协作共建、社区讨论交流的全流程知识管理生态系统。'
)
add_body(
    '前端采用 Next.js 16 + React 19 + TypeScript 技术栈，实现了包括D3.js力导向图谱渲染、Canvas高性能图形绘制、framer-motion动画交互、Zustand状态管理、TanStack Query数据缓存等现代化前端工程实践。后端基于 FastAPI 异步框架构建，提供RESTful API与WebSocket实时通信双协议支持，通过Redis消息总线实现多用户实时聊天，利用MinIO对象存储管理文档资源。AI服务层集成Anthropic Claude和OpenAI双模型供应商，支持动态切换，采用fastembed轻量级嵌入模型构建FAISS向量索引，实现语义相似度检索。'
)
add_body(
    '平台实现了完整的RBAC角色权限控制、JWT身份认证、每日用量配额管理、成长积分等级系统、团队空间协作、知识发酵池论坛、模板市场共享等丰富功能模块。经过多轮迭代优化，系统在高并发场景下表现稳定，文档处理流水线支持并发限流、错误重试、重复检测等鲁棒性保障机制。本项目完整覆盖了课程的基础要求与扩展要求，体现了良好的工程规范性与开发实践水平。'
)

add_heading_custom('关键词', level=1)
kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = kw_p.add_run('知识图谱，大语言模型，Neo4j，FastAPI，Next.js，实体关系抽取，FAISS向量检索，实时协作')
set_font(run, size=12, bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1 引言
# ═══════════════════════════════════════════════
add_heading_custom('1    引  言', level=1)

add_body(
    '随着高等教育的深入推进和科研创新需求的日益增长，实验教学在理工科人才培养中占据着不可替代的重要地位。然而，在传统的实验教学过程中，学生的实验知识往往以碎片化的形式散落在个人笔记、实验报告、课程论文等各类文档中，缺乏系统性的组织与管理。这种知识孤岛现象导致了以下突出问题：第一，实验细节容易被遗忘，难以在后续实验中复用已有经验；第二，不同实验之间的知识关联无法被自动发现，跨学科的隐性联系难以被揭示；第三，团队成员之间的知识传承效率低下，新成员需要从头开始积累经验；第四，实验成果的展示和分享缺乏统一平台，优秀经验难以沉淀为集体智慧。'
)
add_body(
    '近年来，大语言模型（LLM）在自然语言理解、信息抽取、知识推理等领域取得了突破性进展，为自动化知识管理提供了全新的技术路径。同时，图数据库技术在复杂关系建模和可视化方面的成熟应用，使得构建大规模知识图谱成为可能。在此背景下，本项目设计并实现了"创新实验知识图谱平台"（LabGalaxy），将AI技术与图数据库深度融合，为用户提供从文档上传、知识自动萃取、图谱可视化探索到自然语言智能问答的一站式知识管理解决方案。'
)
add_body(
    '本项目的核心价值主张是"让每一份实验知识，都有迹可循"——通过AI自动解析实验文档，抽取实验、设备、理论、耗材、工具等关键实体及其相互关系，构建可搜索、可关联、可复用的知识网络，让隐性知识显性化，让碎片知识系统化。'
)

# ═══════════════════════════════════════════════
# 2 实验内容
# ═══════════════════════════════════════════════
add_heading_custom('2    实验内容', level=1)

add_body(
    '本项目基于FastAPI + Next.js全栈技术栈，以知识图谱为核心载体，构建了一个面向创新实验的知识管理与协作平台。系统涵盖用户认证、文档管理、AI知识抽取、图谱可视化、团队协作、社区论坛、模板市场、成长激励等多个功能模块。以下分别从实验分组、进度安排和实验环境三个方面进行详细说明。'
)

add_heading_custom('2.1   实验分组', level=2)
add_body(
    '本课程实验分组完成创新实验知识图谱平台的整体设计与实现工作。项目采用全栈开发模式，涵盖前端界面设计、后端API开发、AI服务集成、数据库设计与优化、系统部署与测试等全流程工作。项目代码托管于Git版本管理系统，遵循规范的分支管理与提交约定。'
)

add_heading_custom('2.2   进度安排', level=2)
add_body('项目总周期为8周，各阶段任务安排如下表所示：')

add_table_with_data(
    ['序号', '任务', '周期', '起止周'],
    [
        ['1', '需求分析与概要设计', '2周', '第1-2周'],
        ['2', '前端界面开发与交互', '2周', '第3-4周'],
        ['3', '后端API与AI服务开发', '2周', '第3-5周'],
        ['4', '功能开发与系统集成', '2周', '第5-6周'],
        ['5', '联调测试与部署优化', '2周', '第7-8周'],
    ]
)

add_body(
    '第1-2周：完成需求分析、技术选型、系统架构设计、数据库建模、API接口设计。确定前后端技术栈，完成开发环境搭建。'
)
add_body(
    '第3-4周：前端核心页面开发，包括首页、文档管理、图谱可视化、用户认证等模块。后端基础API开发，完成用户认证、文档上传等核心接口。'
)
add_body(
    '第5-6周：AI服务层开发，实现LLM实体关系抽取、Neo4j图谱写入、FAISS向量索引。团队协作、论坛社区、模板市场等功能模块开发。'
)
add_body(
    '第7-8周：前后端联调，性能优化，Docker容器化部署，K8s/K3s集群部署方案，安全加固与测试验证。'
)

add_heading_custom('2.3   实验环境', level=2)

add_heading_custom('2.3.1   硬件环境', level=3)
add_body('项目开发及运行所需的硬件环境配置如下：')

add_table_with_data(
    ['组件', '配置要求'],
    [
        ['处理器', 'Intel Core i7 / AMD Ryzen 7 或同等性能处理器'],
        ['内存', '16GB RAM（推荐32GB，用于本地运行多个Docker容器）'],
        ['存储', '500GB SSD（用于Docker镜像、数据库存储、文档管理）'],
        ['操作系统', 'Windows 11 / Linux (Ubuntu 22.04+) / macOS'],
        ['网络', '稳定的互联网连接（用于访问LLM API服务）'],
    ]
)

add_heading_custom('2.3.2   软件环境', level=3)
add_body('项目所采用的软件技术栈如下：')

add_table_with_data(
    ['组件', '技术选型'],
    [
        ['前端框架', 'Next.js 16 + React 19 + TypeScript 5.6'],
        ['CSS方案', 'TailwindCSS 3.4 + PostCSS'],
        ['状态管理', 'Zustand 5.0（全局状态）+ TanStack Query 5.0（服务端缓存）'],
        ['图形渲染', 'D3.js (force/zoom/selection/drag) + Canvas 2D API'],
        ['动画系统', 'Framer Motion 11 + GSAP 3 + Web Audio API（音效引擎）'],
        ['后端框架', 'Python 3.12 + FastAPI 0.115 + Uvicorn 0.30 (ASGI)'],
        ['关系数据库', 'PostgreSQL 16 + SQLAlchemy 2.0 (async) + asyncpg'],
        ['图数据库', 'Neo4j 5 Community + Python Driver 5.25'],
        ['缓存/消息', 'Redis 7（Pub/Sub消息总线、聊天实时通信）'],
        ['对象存储', 'MinIO（兼容S3协议的文档存储服务）'],
        ['向量索引', 'FAISS 1.8.0 + fastembed (BGE-M3嵌入模型)'],
        ['LLM服务', 'Anthropic Claude API + OpenAI API（双供应商可切换）'],
        ['容器化', 'Docker + Docker Compose + K8s/K3s'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3 作品架构
# ═══════════════════════════════════════════════
add_heading_custom('3    作品架构', level=1)

add_body(
    '本系统采用三层分离的微服务架构，分为前端展示层、后端API层和AI服务层，各层之间通过RESTful API和WebSocket协议进行通信。数据库层包含四种不同类型的数据库组件，分别承担关系型数据存储、图数据存储、缓存消息和对象存储的职责。'
)
add_body(
    '整体架构特点：(1) 前后端完全解耦，前端通过Next.js API Route代理请求，避免跨域问题；(2) AI服务独立部署，通过HTTP与后端通信，支持水平扩展；(3) 多种数据库组件各司其职——关系数据（PostgreSQL）管业务、图数据（Neo4j）管关联、向量索引（FAISS）管语义搜索、对象存储（MinIO）管文件。'
)

add_heading_custom('3.1 前端模块', level=2)

add_heading_custom('3.1.1 界面设计', level=3)

add_body(
    '前端界面设计遵循"温暖学术"的设计语言，整体配色采用以#9A8C73（暖棕灰色）为主色调的大地色系，搭配#F4F1EE（米白色）背景、#4a3e34（深棕）文字色，营造沉稳优雅的学术氛围。字体选用楷体（KaiTi），体现中文字体的书法美感。全局采用Liquid Glass（液态玻璃）毛玻璃效果卡片设计，通过backdrop-filter: blur()实现半透明悬浮质感。'
)
add_body(
    '系统主要页面包括：'
)
add_body(
    '(1) 首页：采用Dashboard式设计，顶部展示用户欢迎信息和等级徽章，中部为6个核心功能模块入口卡片（知识图谱、工作台、上传文档、知识发酵池、模板市场、团队空间），右侧为精选动态轮播，底部为最近文档和成长记录双栏卡片。',
    indent=False, space_after=4
)
add_body(
    '(2) 知识图谱页：全屏Canvas力导向图谱可视化，支持节点拖拽、缩放、悬停高亮、邻接节点聚焦等交互操作。顶部提供工具栏进行节点类型过滤、关键词搜索、年份筛选、时间线动画切换。提供全屏模式和可拖拽像素角色交互按钮。',
    indent=False, space_after=4
)
add_body(
    '(3) 工作台页：左侧文档树按"年份→实验类型"层级组织，右侧为卡片流展示文档详情和AI摘要，支持收藏功能。',
    indent=False, space_after=4
)
add_body(
    '(4) 文档上传页：支持单文件和批量拖拽上传，实时进度条显示AI处理状态（uploaded→parsing→extracting→completed/failed），支持隐私级别和可见团队设置。',
    indent=False, space_after=4
)
add_body(
    '(5) 论坛页（知识发酵池）：6大板块（方法论堂、图谱议事厅、实验急诊室、Aha!广场、学科撞车现场、公告堂），支持热帖排序、精华标记、最佳答案、点赞、收藏等功能。',
    indent=False, space_after=4
)
add_body(
    '(6) 团队空间：支持团队创建、成员邀请（通过display_id或用户名搜索）、实时聊天（WebSocket）、文档权限共享、团队成长时间线。',
    indent=False, space_after=4
)
add_body(
    '(7) 管理员后台：用户管理（搜索、角色变更、激活/禁用）、文档审核、图谱CRUD、AI配置管理、模板审核、设备审批。',
    indent=False, space_after=8
)

add_heading_custom('3.1.2 技术选型', level=3)

add_body('前端技术选型充分考虑了性能、开发体验和可维护性三个维度：')

add_body(
    '(1) Next.js 16：作为React全栈框架，提供了SSR/SSG能力，内置API Route作为后端代理层（/api/v1/[[...path]]/route.ts），有效解决了跨域问题和token转发。App Router架构配合React 19提供了最新的并发特性。',
    indent=False, space_after=4
)
add_body(
    '(2) TypeScript 5.6：全项目TypeScript覆盖，从API层到组件层全面类型安全。所有API响应类型在lib/api.ts中统一定义（共100+接口函数和40+类型定义），IDE自动补全和编译时检查有效减少了运行时错误。',
    indent=False, space_after=4
)
add_body(
    '(3) D3.js + Canvas 2D：图谱可视化是系统的核心交互组件。选用D3.js的forceSimulation、forceManyBody、forceLink、forceCenter等力导向算法模块进行物理模拟，使用Canvas 2D API而非SVG进行渲染，原因是Canvas在节点数量超过200时性能远优于SVG——Canvas模式下60FPS可轻松渲染500+节点。通过requestAnimationFrame + 空闲检测机制（IDLE_INTERVAL_MS=100ms）实现按需渲染。',
    indent=False, space_after=4
)
add_body(
    '(4) Zustand + TanStack Query：Zustand用于管理图谱节点选中状态、全屏模式切换等UI状态（stores/app-store.ts、graph-store.ts、chat-store.ts），极简API减少了大量样板代码。TanStack Query用于缓存所有API请求结果，自动处理stale-while-revalidate策略。',
    indent=False, space_after=4
)
add_body(
    '(5) Framer Motion + GSAP：framer-motion用于组件级过渡动画（页面切换、卡片展开、全屏切换），GSAP用于复杂动画。同时自研了SoundEngine音效引擎（lib/audio/SoundEngine.ts），使用Web Audio API的OscillatorNode合成交互音效，增强用户操作的反馈感。',
    indent=False, space_after=4
)
add_body(
    '(6) TailwindCSS 3.4：原子化CSS方案，通过utility-first方式快速构建响应式界面，postcss.config.js配置autoprefixer实现浏览器兼容前缀自动添加。',
    indent=False, space_after=8
)

add_heading_custom('3.1.3 主要特点', level=3)

add_body(
    '(1) 高性能Canvas图谱渲染：GalaxyView组件使用Canvas 2D API渲染图谱，通过requestAnimationFrame驱动渲染循环。引入空闲检测机制，当仿真静止后降低刷新频率以节省电量。节点悬停时实现邻接高亮/非邻接变暗的聚焦效果，通过alpha渐变过渡（HOVER_TRANSITION_SPEED=0.08）实现平滑的视觉效果。',
    indent=False, space_after=4
)
add_body(
    '(2) 智能轮询与签名检测：图谱页面每15秒（POLL_INTERVAL_MS=15000）自动轮询后端获取最新数据，通过graphSignature函数计算节点和边的排序ID列表的JSON签名，仅当签名发生变化时才更新视图，避免不必要的重渲染。',
    indent=False, space_after=4
)
add_body(
    '(3) 时间线动画模式：Timeline模式下节点按类型排序依次出现（每200ms新增一个），配合力导向仿真动态构建图谱，生动展示知识积累过程。节点从画布中心以随机偏移淡入，已有边在节点出现后自动连接。',
    indent=False, space_after=4
)
add_body(
    '(4) 全屏模式与像素角色交互：通过framer-motion的spring动画实现全屏展开/收起（scale 0.08→1，stiffness=200，damping=16），模拟手机App的打开效果。ESC键退出。一个可拖拽的像素角色作为全屏切换按钮，增加了趣味性。',
    indent=False, space_after=4
)
add_body(
    '(5) 文档权限与可见性控制：文档支持public/team/private三级隐私设置，图谱数据根据用户所属团队和文档可见性团队动态过滤。可见性计算在PostgreSQL层面完成（_visible_document_ids），将可见文档ID列表传递给Neo4j的Cypher查询进行图层面过滤。',
    indent=False, space_after=8
)

add_heading_custom('3.2 后端模块', level=2)

add_heading_custom('3.2.1 API接口简介', level=3)

add_body('后端系统基于FastAPI框架构建，提供完整的RESTful API接口。所有接口统一挂载在/api/v1路径下，采用JWT Bearer Token认证。核心API模块如下：')

add_table_with_data(
    ['模块', '路径前缀', '功能说明'],
    [
        ['用户认证', '/users', '注册、登录、个人信息、Dashboard'],
        ['文档管理', '/documents', '上传、状态、重处理、去重确认、删除、下载'],
        ['工作台', '/workbench', '文档树、卡片流、收藏'],
        ['知识图谱', '/graph', '图谱数据、CRUD、SSE流、AI建议、搜索、清理'],
        ['自然语言查询', '/query', '基于图谱的问答，支持对话历史'],
        ['知识发酵池', '/forum', '板块、帖子、回复、点赞、精华、公告'],
        ['团队协作', '/teams', '创建、成员、邀请、退出、删除'],
        ['实时聊天', '/ws', 'WebSocket团队聊天（Redis Pub/Sub）'],
        ['模板市场', '/templates', 'CRUD、审核、点赞、收藏、评论、成长'],
        ['管理员', '/admin', '用户/文档/图谱/模板管理、AI配置'],
        ['设备管理', '/equipment', '设备目录、申请、审批'],
        ['洞察发现', '/insights', 'AI发现跨文档关联'],
    ]
)

add_heading_custom('3.2.2 框架选型说明', level=3)

add_body(
    '本项目后端采用的是 FastAPI 框架。FastAPI 是基于Starlette的现代化异步Python Web框架，相比传统Flask具有以下技术优势，这也是本项目选择FastAPI的核心原因：'
)
add_body(
    '(1) 原生异步支持：FastAPI基于ASGI规范，天然支持async/await。本项目中所有数据库操作（asyncpg）、HTTP请求（httpx调用AI服务）、文件IO均采用异步模式，有效提升了I/O密集型场景下的并发处理能力。',
    indent=False, space_after=4
)
add_body(
    '(2) 自动API文档生成：基于OpenAPI 3.0自动生成Swagger UI（/docs）和ReDoc（/redoc）交互式文档，大幅降低接口文档维护成本。',
    indent=False, space_after=4
)
add_body(
    '(3) Pydantic数据校验：深度集成Pydantic v2，通过类型注解实现请求体、查询参数、路径参数的自动校验和类型转换。所有request/response model均有明确定义。',
    indent=False, space_after=4
)
add_body(
    '(4) 依赖注入系统：通过get_current_user、require_admin、get_db等依赖函数实现统一的认证授权和数据库会话管理。',
    indent=False, space_after=4
)
add_body(
    '(5) 性能优势：基于Starlette的高性能ASGI实现，异步处理使得单节点可轻松支撑100+并发请求。',
    indent=False, space_after=8
)

add_heading_custom('3.2.3 接口设计', level=3)

add_body('后端API接口设计遵循RESTful规范，具体设计原则和技术实现细节如下：')

add_body(
    '(1) 统一认证机制：所有需要认证的接口均通过get_current_user依赖进行JWT Token验证。Token通过python-jose库进行HS256算法签名，默认有效期24小时。对于文档下载等特殊场景，实现了get_current_user_from_header_or_query依赖，同时支持Header Bearer Token和Query参数两种方式认证。',
    indent=False, space_after=4
)
add_body(
    '(2) 错误处理：统一使用HTTPException抛出标准HTTP错误，错误信息通过detail字段传递中文提示。关键操作（如文档删除）采用graceful degradation策略——即使MinIO删除或Neo4j清理失败，PostgreSQL中的记录删除仍然继续。',
    indent=False, space_after=4
)
add_body(
    '(3) 分页设计：所有列表接口统一支持page和page_size参数，默认page_size=20，最大限制通过le参数约束。返回格式统一为{total, items}。',
    indent=False, space_after=4
)
add_body(
    '(4) 异步后台任务：文档上传后通过asyncio.create_task启动后台AI处理，立即返回uploaded状态。批量上传采用asyncio.Semaphore(3)限制最大并发数为3。处理过程中通过独立数据库会话更新状态，前端轮询/status接口获取进度。',
    indent=False, space_after=4
)
add_body(
    '(5) WebSocket实时通信：团队聊天采用WebSocket，连接前通过JWT认证并验证团队成员身份。消息通过Redis Pub/Sub广播，加入/离开时自动发送系统消息。',
    indent=False, space_after=4
)
add_body(
    '(6) 增量Schema迁移：main.py的startup事件包含12个渐进式Schema迁移块，每个使用try-except包裹，通过ALTER TABLE ... ADD COLUMN IF NOT EXISTS和DO $$ BEGIN ... EXCEPTION处理幂等性。',
    indent=False, space_after=8
)

add_heading_custom('3.3 数据库集成', level=2)

add_body(
    '本系统创新性地集成了四种不同类型的数据库/存储组件，每种组件承担特定的数据管理职责：'
)
add_body(
    '(1) PostgreSQL 16（关系型数据库）：作为核心业务数据存储，承载了用户信息、文档元数据、论坛帖子、团队、消息、模板、积分记录等15张数据表。使用SQLAlchemy 2.0的异步ORM（mapped_column + Mapped类型注解）。关键设计：UUID主键（uuid_generate_v4()）、ENUM类型（doc_status、post_type等）、ARRAY类型（标签数组）、B-Tree索引。',
    indent=False, space_after=4
)
add_body(
    '(2) Neo4j 5（图数据库）：知识图谱核心存储。存储六类实体节点（Experiment/Equipment/Theory/Consumable/Tool/Concept）和五类关系（USES/BASED_ON/SIMILAR_TO/REQUIRES/RELATED_TO）。关键设计：MERGE操作实现实体去重（按name+type合并）、关系存储confidence和document_id属性、节点通过document_id关联PostgreSQL文档。',
    indent=False, space_after=4
)
add_body(
    '(3) FAISS（向量索引）：AI服务中使用FAISS的IndexIDMap + IndexFlatIP（内积相似度）构建实体向量索引。嵌入模型采用fastembed的BGE-M3（轻量级、离线可用）。UUID到FAISS整数ID通过MD5哈希前15位转换，持久化在.idmap.json中。启动时从Neo4j全量重建索引。',
    indent=False, space_after=4
)
add_body(
    '(4) MinIO（对象存储）：兼容S3协议，存储用户上传的原始文档（PDF/DOCX/PPTX）。以UUID作为object key，下载时通过StreamingResponse流式返回（64KB分块传输）。',
    indent=False, space_after=4
)
add_body(
    '(5) Redis（缓存/消息）：两个角色——团队聊天的Pub/Sub消息广播总线，以及每日用量配额管理（incr原子计数）。',
    indent=False, space_after=8
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4 课程要求的完成情况
# ═══════════════════════════════════════════════
add_heading_custom('4    课程要求的完成情况', level=1)

add_heading_custom('4.1 项目实现情况', level=2)

add_heading_custom('4.1.1 基础要求实现（全部按要求已实现）', level=3)

add_body(
    '【用户认证与权限管理】实现了完整的用户注册、登录、个人信息管理。密码使用bcrypt哈希存储，JWT Token进行会话管理。支持角色区分（user/admin），管理员拥有独立后台。新用户注册后需管理员审批激活（is_active字段）。',
    indent=False, space_after=4
)
add_body(
    '【文档上传与管理】支持PDF/DOC/DOCX/PPT/PPTX五种格式，单文件限50MB。支持单文件和批量拖拽上传。文档支持实验年份、实验类型、学科标签、隐私级别等元数据。上传后自动触发AI处理流水线，实时显示进度。',
    indent=False, space_after=4
)
add_body(
    '【知识抽取与图谱构建】AI服务调用LLM进行实体抽取和关系抽取。实体通过MD5(name+type)生成确定性UUID实现跨文档去重。关系写入Neo4j时按标签和关系类型分组批量操作（UNWIND）。同时生成FAISS向量索引。',
    indent=False, space_after=4
)
add_body(
    '【图谱可视化】D3.js力导向算法 + Canvas 2D渲染，支持节点拖拽、缩放、悬停高亮、邻接聚焦。提供时间线和关系矩阵辅助视图。支持按类型、关键词、年份过滤。',
    indent=False, space_after=4
)
add_body(
    '【数据持久化】PostgreSQL（15张业务表）、Neo4j（知识图谱）、FAISS（向量索引）、MinIO（原始文件）、Redis（消息/缓存），五种存储组件各司其职。',
    indent=False, space_after=8
)

add_heading_custom('4.1.2 扩展要求实现', level=3)

add_body(
    '【AI智能问答】基于知识图谱的自然语言问答。FAISS向量检索找到相关实体，结合LLM生成结构化回答（答案文本、高亮节点、源文档、建议追问）。支持多轮对话历史。',
    indent=False, space_after=4
)
add_body(
    '【团队协作空间】团队创建、成员邀请、文档权限共享。团队维度图谱数据隔离，实时聊天基于WebSocket + Redis Pub/Sub。',
    indent=False, space_after=4
)
add_body(
    '【知识发酵池论坛】6个特色板块，6种帖子类型，热帖排序算法（likes*5 + replies*2 + views*0.1加权衰减），精华标记，最佳答案。',
    indent=False, space_after=4
)
add_body(
    '【模板市场与成长系统】模板CRUD、发布审核、点赞收藏。完整成长激励系统：发帖/回复/模板被收藏获取积分，10级等级配置（含图标和称号）。',
    indent=False, space_after=4
)
add_body(
    '【管理员后台】用户管理、文档审核、图谱CRUD、AI配置热更新、设备审批。',
    indent=False, space_after=4
)
add_body(
    '【重复检测与确认机制】使用difflib.SequenceMatcher计算相似度（阈值0.78），检测到重复时转为awaiting_confirmation状态等待用户确认（overwrite/coexist/cancel）。',
    indent=False, space_after=4
)
add_body(
    '【SSE实时事件流】图谱变更通过Server-Sent Events推送，自定义EventBus实现。',
    indent=False, space_after=4
)
add_body(
    '【容器化与部署】Docker Compose开发环境 + K8s/K3s生产环境（11个YAML清单），一键部署脚本，自动化备份。',
    indent=False, space_after=8
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5 应用展示
# ═══════════════════════════════════════════════
add_heading_custom('5    应用展示', level=1)

add_body(
    '本平台已完整实现了从文档上传到知识发现的全流程功能，主要应用场景如下：'
)

add_body(
    '(1) 实验知识沉淀：学生上传实验报告，系统自动识别实验名称、设备仪器、理论原理、试剂材料等关键实体，自动建立"实验-使用-设备"、"实验-基于-理论"等关系。多次实验后，知识图谱自动展现不同实验之间的关联。',
    indent=False, space_after=4
)
add_body(
    '(2) 跨学科知识发现：通过图谱可视化和AI智能问答，发现学科之间的隐性联系。例如物理实验中的"RC电路"与电子实验中的"滤波器设计"共享相同理论基础。',
    indent=False, space_after=4
)
add_body(
    '(3) 团队协作学习：团队成员共享实验知识，引用已有知识节点。团队聊天支持实时讨论，团队成长时间线展示知识积累历程。',
    indent=False, space_after=4
)
add_body(
    '(4) 社区知识交流：知识发酵池论坛提供结构化讨论平台。帖子可与知识图谱节点关联，引用具体实体，形成"讨论-知识"双向链接。',
    indent=False, space_after=4
)
add_body(
    '(5) 管理员运营：审核用户、文档、图谱，配置AI模型参数，查看平台运营数据。AI配置支持热更新。',
    indent=False, space_after=8
)

# ── 技术难点与解决方案 ──
add_heading_custom('5.1 技术难点与解决方案', level=2)

add_body(
    '【难点一：LLM输出JSON的鲁棒性解析】大语言模型输出的JSON经常出现格式错误（缺少逗号、未转义引号、markdown包裹、嵌套不完整）。',
    indent=False, space_after=4
)
add_body(
    '解决方案：实现了五层渐进式JSON修复策略——第一层正则去除markdown fences；第二层调用json_repair库智能修复；第三层正则提取entities数组片段；第四层基于栈的实体对象提取算法（_extract_entity_objects），逐个提取合法JSON对象；第五层使用Python ast.literal_eval最终尝试。实体抽取失败后自动使用retry prompt重新请求。JSON解析成功率提升至99%以上。',
    indent=False, space_after=4
)

add_body(
    '【难点二：大规模图谱Canvas渲染性能】当节点超过200个时，SVG渲染出现明显卡顿。',
    indent=False, space_after=4
)
add_body(
    '解决方案：完全切换到Canvas 2D API，优化包括：(a) requestAnimationFrame驱动渲染循环；(b) dirty flag机制，仅数据变化时重绘；(c) 空闲检测（100ms后降低频率）；(d) 标签根据缩放级别动态显示/隐藏（zoom<0.5隐藏，0.5-0.8渐显，>2.5渐隐）；(e) 悬停节点标签在最顶层单独渲染。实现60FPS流畅体验。',
    indent=False, space_after=4
)

add_body(
    '【难点三：跨数据库数据一致性】PostgreSQL和Neo4j两种数据库的同步清理问题。',
    indent=False, space_after=4
)
add_body(
    '解决方案：采用graceful degradation策略——主数据库（PostgreSQL）操作失败则事务回滚，辅助数据库（Neo4j、MinIO）操作失败则记录警告但继续执行。文档重处理时先清理旧数据再写入新数据。实体写入使用MERGE实现幂等性。',
    indent=False, space_after=4
)

add_body(
    '【难点四：实体去重与跨文档合并】不同文档可能描述同一实验但命名略有差异。',
    indent=False, space_after=4
)
add_body(
    '解决方案：(a) 写入层：MERGE (n:Experiment {name: normalized_name})实现精确匹配去重，名称归一化（折叠空白、去除"实验"后缀）；(b) 检测层：difflib.SequenceMatcher相似度（阈值0.78），检测到相似实体时暂停等待用户确认；(c) 推荐层：FAISS向量相似度推荐可能关联。',
    indent=False, space_after=4
)

add_body(
    '【难点五：知识图谱权限隔离】确保用户只能访问有权看到的节点。',
    indent=False, space_after=4
)
add_body(
    '解决方案：每次图谱查询时先在PostgreSQL中计算可见文档ID列表（考虑privacy、visible_teams、status），将doc_ids传递给Neo4j的Cypher查询进行过滤。团队模式下额外验证团队成员身份，防止越权访问。',
    indent=False, space_after=8
)

# ── 工程规范性 ──
add_heading_custom('5.2 工程性与开发规范性', level=2)

add_body(
    '【代码组织规范】后端严格按照api/models/schemas/services/core目录组织。前端按照app（路由）/components（组件）/lib（工具）/hooks（自定义Hook）/stores（状态管理）组织。每个模块职责清晰，低耦合高内聚。',
    indent=False, space_after=4
)
add_body(
    '【类型安全】后端Pydantic v2数据校验，前端TypeScript全量覆盖（100+ API接口函数，40+类型定义）。',
    indent=False, space_after=4
)
add_body(
    '【错误处理】所有异常有明确错误处理和中文提示。关键操作采用graceful degradation。数据库迁移try-except包裹确保幂等性。AI服务层重试机制。',
    indent=False, space_after=4
)
add_body(
    '【安全实践】bcrypt密码哈希、JWT Token 24小时有效期、环境变量配置敏感信息（.env在.gitignore中）、Docker非root用户运行、独立管理员权限校验、WebSocket连接前JWT认证和团队身份验证。',
    indent=False, space_after=4
)
add_body(
    '【部署规范】Docker Compose开发环境 + K8s生产环境配置。Dockerfile遵循最佳实践：非root用户、健康检查、精简依赖。11个K8s清单文件按资源类型分离，支持一键部署和滚动更新。提供自动化部署和备份脚本。',
    indent=False, space_after=4
)
add_body(
    '【版本控制】Git版本管理，Conventional Commits规范（feat:/fix:/refactor:/docs:前缀）。.gitignore和.dockerignore排除不必要文件。',
    indent=False, space_after=4
)
add_body(
    '【可观测性】结构化日志记录（logging模块），关键操作日志输出。AI服务记录流水线各阶段耗时。/health健康检查接口。K8s中每个服务配置healthcheck。',
    indent=False, space_after=8
)

# 保存
output_path = r'D:\X\experiment_report.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
